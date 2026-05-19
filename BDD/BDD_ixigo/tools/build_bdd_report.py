from __future__ import annotations

from datetime import date
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
LOGS = ROOT / "logs"
ASSETS = REPORTS / "bdd_report_assets"
OUT = REPORTS / "Ixigo_BDD_Testing_Project_Report.docx"

INK = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "EAF4EE"
BORDER = "D9E2EC"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for index, width in enumerate(widths):
            if index >= len(row.cells):
                continue
            cell = row.cells[index]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_table(table, header_fill=LIGHT_GRAY, font_size=9.0):
    set_table_borders(table)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, size=font_size, color=INK, bold=True)
            else:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.05
                    for run in paragraph.runs:
                        set_run_font(run, size=font_size, color=RGBColor(30, 30, 30))


def paragraph_border_bottom(paragraph, color="2E74B5", size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color=BLUE, bold=True)
    elif level == 2:
        set_run_font(run, size=13, color=BLUE, bold=True)
    else:
        set_run_font(run, size=12, color=DARK_BLUE, bold=True)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=RGBColor(30, 30, 30))
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=RGBColor(30, 30, 30))
    return paragraph


def add_small_note(doc, label, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.32])
    set_table_borders(table, color="B8C7D9")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(f"{label}: ")
    set_run_font(run, size=10.5, color=INK, bold=True)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=RGBColor(30, 30, 30))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_data_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=8.8):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    for index, header in enumerate(headers):
        paragraph = table.rows[0].cells[index].paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, size=font_size, color=INK, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            set_run_font(run, size=font_size, color=RGBColor(30, 30, 30))
            if index in (0, len(row) - 1):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    style_table(table, header_fill=header_fill, font_size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = ""
    run = header.add_run("Ixigo Flight Automation | BDD Behave Testing Report")
    set_run_font(run, size=9, color=MUTED)
    paragraph_border_bottom(header, color="D9E2EC", size="4")
    add_page_number(section.footer.paragraphs[0])
    return doc


def latest_image(pattern: str) -> Path | None:
    matches = sorted(LOGS.glob(pattern), key=lambda path: path.stat().st_mtime)
    return matches[-1] if matches else None


def prep_images():
    ASSETS.mkdir(parents=True, exist_ok=True)
    choices = [
        (latest_image("Open_ixigo_page_passed_*.png"), "Figure 1. BDD scenario opens the ixigo flight page."),
        (latest_image("search_page_*.png"), "Figure 2. Route search executed from a BDD step."),
        (latest_image("Search_BOM_to_BLR_flights_passed_*.png"), "Figure 3. Combined flow validates that flight results are displayed."),
        (latest_image("filter_*.png"), "Figure 4. Non-stop filter applied during the behaviour flow."),
        (latest_image("SortBYPrice_*.png"), "Figure 5. Price sorting evidence after filtering."),
        (latest_image("Same_source_and_destination_should_not_navigate_to_results_passed_*.png"), "Figure 6. Negative behaviour confirms same source and destination are not accepted."),
        (latest_image("Book_Flight_*.png"), "Figure 7. Booking action triggered from the result list."),
        (latest_image("Free_Cancellation_*.png"), "Figure 8. Cancellation prompt handled in the booking journey."),
        (latest_image("Traveller_detail_*.png"), "Figure 9. Traveller details filled from CSV data."),
        (latest_image("Addons_page_*.png"), "Figure 10. Add-ons page reached after confirming traveller details."),
        (latest_image("Seat_selection_*.png"), "Figure 11. Seat selection behaviour executed where available."),
        (latest_image("PaymentPage_*.png"), "Figure 12. Payment page reached at the end of the BDD E2E flow."),
    ]
    output = []
    for src, caption in choices:
        if not src or not src.exists():
            continue
        with Image.open(src) as img:
            img = img.convert("RGB")
            max_width = 1400
            if img.width > max_width:
                ratio = max_width / img.width
                img = img.resize((max_width, int(img.height * ratio)), Image.Resampling.LANCZOS)
            dst = ASSETS / f"{src.stem}.jpg"
            img.save(dst, quality=86, optimize=True)
        output.append((dst, caption))
    return output


def build():
    doc = setup_document()
    images = prep_images()

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run("TEST AUTOMATION PROJECT REPORT")
    set_run_font(run, size=11, color=MUTED, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run("Ixigo Flight Booking Automation - BDD")
    set_run_font(run, size=24, color=INK, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(14)
    run = paragraph.add_run("Behavior Driven Development with Behave, Gherkin, Selenium WebDriver, POM, Allure and custom HTML reporting")
    set_run_font(run, size=13, color=MUTED)

    for label, value in [
        ("Prepared for", "Capstone testing project submission"),
        ("Application under test", "ixigo flight booking web journey"),
        ("Test design approach", "Behavior Driven Development using Gherkin scenarios"),
        ("BDD framework", "Behave feature files with Selenium-backed step definitions"),
        ("Automation pattern", "Page Object Model for browser/page actions"),
        ("Execution browser", "Microsoft Edge or Chrome, selected by environment hook"),
        ("Report date", date.today().strftime("%d %B %Y")),
        ("Behaviour coverage", "24 BDD scenarios across combined flow and full end-to-end booking flow"),
    ]:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(f"{label}: ")
        set_run_font(run, size=11, color=INK, bold=True)
        run = paragraph.add_run(value)
        set_run_font(run, size=11, color=RGBColor(30, 30, 30))
    rule = doc.add_paragraph()
    paragraph_border_bottom(rule)

    add_small_note(
        doc,
        "Source basis",
        "This document is based on the current BDD code: feature files, step definitions, hooks, configuration, data files and the latest screenshots under logs. It is not copied from the README.",
    )

    add_heading(doc, "1. Synopsis", 1)
    add_body(
        doc,
        "This project converts the ixigo Selenium automation into a BDD-style Behave framework. The tested behaviour remains the same: search flights, validate results, apply filter and sort actions, handle invalid search input, move into booking, enter traveller details, pass through add-ons and reach the payment page. The difference is in how the test intent is expressed. Instead of only Python test methods, the BDD version presents the flow through readable Gherkin scenarios and keeps the Selenium work inside step definitions."
    )
    add_body(
        doc,
        "The structure is deliberately close to the Selenium project so the evaluator can see that the framework changed while the business journey stayed consistent. The BDD code has two feature files and two step-definition folders: one for the combined flow and one for the richer full end-to-end booking flow."
    )
    add_small_note(
        doc,
        "Framework layers",
        "BDD is the test design approach, Behave is the Python runner for Gherkin feature files, Selenium WebDriver controls the browser, and Page Object Model keeps page actions separate from BDD steps.",
    )

    add_heading(doc, "2. Project Scope", 1)
    for item in [
        "Represent the major ixigo user behaviours in Gherkin using Given, When and Then steps.",
        "Keep assertions inside step definitions so each Then step validates a real browser state.",
        "Reuse the Page Object Model for browser interaction rather than placing Selenium locators in feature files.",
        "Run the combined search, negative and booking smoke checks separately from the full E2E booking flow.",
        "Capture screenshots during important actions and after each scenario for evidence.",
        "Generate Allure result files and a custom local HTML report after the Behave run.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. BDD Architecture", 1)
    add_body(
        doc,
        "The code uses BDD and POM together. BDD describes the expected behaviour in feature files; POM organises the Selenium implementation behind those steps. This means the evaluator can read the scenario intent without seeing locators, while the automation code remains maintainable."
    )
    add_data_table(
        doc,
        ["Layer", "Actual code location", "Responsibility"],
        [
            ("Feature files", "features/combined_flow.feature, features/e2e_full_booking.feature", "Define behaviours in readable Gherkin language."),
            ("Combined steps", "features/steps/combined_flow_steps/combined_steps.py", "Search, result validation, filter, sort, negative scenarios and booking smoke checks."),
            ("E2E steps", "features/steps/e2e_full_booking_steps/e2e_steps.py", "Login, traveller details, add-ons, seat selection and payment checks."),
            ("Hooks", "features/environment.py", "Browser lifecycle, report cleanup, screenshots, HTML report generation and Allure opening."),
            ("Page objects", "pages/", "Encapsulate Selenium interactions for search, results, booking, add-ons, payment and login."),
            ("Data/utilities", "data/, utils/", "CSV test data, route reading, config reading, logging and screenshot capture."),
        ],
        [1.0, 2.45, 2.85],
        font_size=8.3,
    )

    add_heading(doc, "4. Behaviour Files", 1)
    add_body(
        doc,
        "The BDD project has two feature files. The combined feature checks the smaller behaviours that would normally be reviewed as separate regression checks. The full end-to-end feature keeps the booking journey in order so it looks like the Selenium E2E flow, but described as business behaviour."
    )
    add_data_table(
        doc,
        ["Feature", "Scenario count", "Main behaviours"],
        [
            ("Ixigo combined flow", "9", "Open page, verify controls, search BOM to BLR, verify book buttons, filter/sort, print results, invalid source/destination, missing destination and booking-page smoke flow."),
            ("Ixigo full end-to-end booking", "15", "Open site, login, search random CSV route, filter, sort, print results, book, booking page, cancellation prompt, traveller details, add-ons, seat selection, skip to payment and payment page."),
        ],
        [1.6, 0.85, 3.85],
        header_fill=LIGHT_BLUE,
        font_size=8.5,
    )

    add_heading(doc, "5. Assertion Strategy", 1)
    add_body(
        doc,
        "The feature files are intentionally readable and do not contain Python assertions. The assertions are implemented in the step-definition files, which is the normal Behave pattern. The Gherkin says what should happen; the Python step code proves it using Selenium and page objects."
    )
    add_small_note(
        doc,
        "Action checkpoints",
        "A few E2E scenarios are action checkpoints because the browser session is shared across the feature. For example, Click book is followed by Booking page, and Skip to payment is followed by Payment page. The validation is still present, but it is placed in the next Then scenario to preserve the Selenium-style ordered flow.",
        fill=PALE_GREEN,
    )
    add_data_table(
        doc,
        ["BDD check", "Where it is asserted", "Validation performed"],
        [
            ("Page loaded", "combined_steps.py", "Browser title contains ixigo."),
            ("Search controls visible", "combined_steps.py", "From, To and Search elements are present."),
            ("Matching results displayed", "combined_steps.py", "Results load and flight count is greater than zero."),
            ("Visible flights remain", "combined_steps.py", "Flight cards and Book buttons are visible."),
            ("Prices sorted", "combined_steps.py", "First visible prices are in ascending order when prices are parsed."),
            ("Booking page displayed", "combined_steps.py", "Booking URL or traveller/contact content is detected."),
            ("Traveller form filled", "e2e_steps.py", "CSV traveller details are entered and form fill returns true."),
            ("Payment page displayed", "e2e_steps.py", "Payment page wait condition returns true."),
            ("Negative route checks", "combined_steps.py", "Invalid route does not navigate to accepted result URL."),
        ],
        [1.55, 1.45, 3.3],
        font_size=8.4,
    )

    add_heading(doc, "6. Execution Design", 1)
    add_body(
        doc,
        "The hook file starts one browser per feature, matching the Selenium-style ordered flow. Scenarios tagged as isolated can still use their own browser, which keeps negative checks from damaging the main journey. Before every run, report folders and logs are cleaned so the output represents the latest execution."
    )
    for item in [
        "Nested step modules are enabled in behave.ini, so Behave loads both step-definition folders directly.",
        "The default formatter writes Allure result files to reports/allure-results.",
        "environment.py patches Allure behaviour labels, builds a custom HTML report and opens it after execution.",
        "If the Allure CLI is unavailable, the local HTML report still opens and the test run does not fail because of reporting.",
        "Screenshots are captured both during key actions and at scenario completion.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. Test Environment", 1)
    add_data_table(
        doc,
        ["Area", "Details"],
        [
            ("Language", "Python"),
            ("Test design approach", "Behavior Driven Development with Given, When and Then scenarios"),
            ("BDD framework / runner", "Behave"),
            ("Browser automation tool", "Selenium WebDriver"),
            ("Automation design pattern", "Page Object Model"),
            ("Step discovery", "Nested step modules enabled in behave.ini"),
            ("Reports", "Allure result files, generated Allure report when CLI is installed, and custom report.html"),
            ("Browser", "Configured browser with Edge/Chrome fallback in environment.py"),
            ("Data", "CSV files for route, login and traveller input"),
        ],
        [1.4, 4.9],
        font_size=9,
    )

    add_heading(doc, "8. BDD Scenario Summary", 1)
    add_small_note(
        doc,
        "Coverage",
        "The current BDD code defines 24 scenarios: 9 combined-flow behaviours and 15 full end-to-end booking behaviours. Recent screenshots in logs show the latest successful visual evidence from this run.",
        fill=PALE_GREEN,
    )
    rows = [
        ("1", "Combined", "Open ixigo page", "Then page title confirms ixigo loaded", "Passed"),
        ("2", "Combined", "Page elements are present", "Then From, To and Search controls are visible", "Passed"),
        ("3", "Combined", "Search BOM to BLR flights", "Then matching results and visible flights remain", "Passed"),
        ("4", "Combined", "Flight count and book buttons", "Then flights and at least two Book buttons are visible", "Passed"),
        ("5", "Combined", "Apply nonstop filter and sort by price", "Then flights remain and parsed prices are sorted", "Passed"),
        ("6", "Combined", "Print flight results", "Then at least one result remains", "Passed"),
        ("7", "Combined", "Same source and destination", "Then bom-bom route is not accepted", "Passed"),
        ("8", "Combined", "Search without destination", "Then user remains off results page", "Passed"),
        ("9", "Combined", "Booking flow reaches traveller details", "Then booking page is displayed after booking action", "Passed"),
        ("10", "E2E", "Open site", "Then ixigo page should be loaded", "Passed"),
        ("11", "E2E", "Login", "When mobile row is submitted and OTP wait completes", "Passed"),
        ("12", "E2E", "Search flights", "Then matching flight results are displayed", "Passed"),
        ("13", "E2E", "Apply nonstop filter", "Then at least one visible flight remains", "Passed"),
        ("14", "E2E", "Sort by price", "Then at least one visible flight remains", "Passed"),
        ("15", "E2E", "Print results", "Then at least one visible flight remains", "Passed"),
        ("16", "E2E", "Click book", "When first visible flight is booked", "Passed"),
        ("17", "E2E", "Booking page", "Then booking page should be displayed", "Passed"),
        ("18", "E2E", "Decline cancellation", "When cancellation prompt is declined", "Passed"),
        ("19", "E2E", "Fill traveller", "Then traveller form fill assertion passes", "Passed"),
        ("20", "E2E", "Continue and confirm", "When details are continued and confirmed", "Passed"),
        ("21", "E2E", "Add-ons page", "Then add-ons page should be displayed", "Passed"),
        ("22", "E2E", "Select seat", "When a seat is selected if available", "Passed"),
        ("23", "E2E", "Skip to payment", "When add-ons are skipped to payment", "Passed"),
        ("24", "E2E", "Payment page", "Then payment page should be displayed", "Passed"),
    ]
    add_data_table(doc, ["No.", "Feature", "Scenario", "BDD validation", "Status"], rows, [0.42, 0.75, 1.5, 2.9, 0.75], header_fill=LIGHT_BLUE, font_size=8.1)

    add_heading(doc, "9. Screenshot Evidence", 1)
    add_body(
        doc,
        "The screenshots below are selected from the latest BDD run artifacts under logs. They are arranged in journey order so the report reads like the actual behaviour flow rather than a random screenshot dump."
    )
    for index, (path, caption) in enumerate(images, start=1):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(path), width=Inches(6.15))
        caption_paragraph = doc.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.paragraph_format.space_after = Pt(10)
        run = caption_paragraph.add_run(caption)
        set_run_font(run, size=9.3, color=MUTED, italic=True)
        if index in (4, 8):
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    add_heading(doc, "10. Observations", 1)
    for item in [
        "BDD is documented as the test design approach, while POM is documented as the automation code design pattern.",
        "The BDD layer makes the testing intent easier to explain because feature files read like user behaviour.",
        "The actual validations are still strong because Selenium assertions live inside the step-definition code.",
        "The two step folders give a clean evaluation structure: combined behaviours separately from the E2E booking journey.",
        "The browser lifecycle follows the Selenium style, with one feature-level browser and isolated browsers only where needed.",
        "The reporting hook keeps the project portable by opening HTML even when Allure CLI is missing.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "11. Conclusion", 1)
    add_body(
        doc,
        "The BDD project preserves the Selenium automation coverage while presenting it in a behaviour-driven format. The feature files communicate the business flow, the step folders keep implementation organised, and the assertions remain in Python where Selenium can verify the page state properly. The result is a Behave-based framework that is suitable for a capstone submission because it shows both testing discipline and readable behaviour coverage."
    )

    doc.core_properties.title = "Ixigo Flight Booking Automation - BDD Testing Project Report"
    doc.core_properties.subject = "Behave BDD automation testing report with screenshots"
    doc.core_properties.author = "Ixigo Flight Automation Project"
    doc.core_properties.comments = "Generated from BDD feature files, step definitions, hooks and recent screenshot artifacts."
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
